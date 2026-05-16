"""
python manage.py create_sample_templates

Створює 3 базові Word шаблони:
- Packing List EN
- Proforma Invoice EN
- CN23 Customs Declaration EN
"""
from django.core.management.base import BaseCommand
from io import BytesIO


class Command(BaseCommand):
    help = 'Створити 3 базові Word шаблони документів'

    def handle(self, *args, **options):
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from django.core.files.base import ContentFile
        from documents.models import DocumentTemplate

        A = WD_ALIGN_PARAGRAPH

        def save(doc, name, doc_type):
            if DocumentTemplate.objects.filter(name=name).exists():
                self.stdout.write(f'  Skip: {name}')
                return
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            tpl = DocumentTemplate(
                name=name, doc_type=doc_type,
                module='sales', language='en',
                is_active=True, is_default=True)
            tpl.template_file.save(
                f'{doc_type}_sample_en.docx',
                ContentFile(buf.getvalue()), save=True)
            self.stdout.write(f'  OK: {name}')

        # ── Packing List ──────────────────────────────────────────────────────
        doc = Document()
        h = doc.add_heading('PACKING LIST', 0)
        h.alignment = A.CENTER
        doc.add_paragraph('')

        for label, var in [
            ('Order №:', '{{order_number}}'),
            ('Date:', '{{order_date}}'),
            ('Tracking:', '{{tracking_number}}'),
            ('Carrier:', '{{carrier_name}}'),
        ]:
            p = doc.add_paragraph()
            p.add_run(f'{label} ').bold = True
            p.add_run(var)

        doc.add_paragraph('')
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        tbl.rows[0].cells[0].text = (
            'FROM:\n{{shipper_name}}\n{{shipper_address}}\n'
            '{{shipper_city}}, {{shipper_country}}\n{{shipper_email}}')
        tbl.rows[0].cells[1].text = (
            'TO:\n{{customer_name}}\n{{customer_address}}\n'
            '{{customer_city}}, {{customer_country}}\n{{customer_email}}')

        doc.add_paragraph('')
        doc.add_heading('Items', level=2)
        tbl2 = doc.add_table(rows=2, cols=6)
        tbl2.style = 'Table Grid'
        for i, h in enumerate(['SKU', 'Description', 'Qty', 'Unit Price', 'Total', 'Weight']):
            c = tbl2.rows[0].cells[i]
            c.text = h
            c.paragraphs[0].runs[0].bold = True
        r = tbl2.rows[1]
        r.cells[0].text = '{%tr for item in items %}{{item.sku}}'
        r.cells[1].text = '{{item.name}}'
        r.cells[2].text = '{{item.quantity}}'
        r.cells[3].text = '{{item.unit_price}} {{currency}}'
        r.cells[4].text = '{{item.total_price}} {{currency}}'
        r.cells[5].text = '{{item.weight}} kg{%tr endfor %}'

        doc.add_paragraph('')
        for label, var in [
            ('Total weight:', '{{total_weight}} kg'),
            ('Total amount:', '{{total_amount}} {{currency}}'),
        ]:
            p = doc.add_paragraph()
            p.add_run(f'{label} ').bold = True
            p.add_run(var)
        doc.add_paragraph('{{notes}}')
        doc.add_paragraph('Generated: {{generated_date}}')
        save(doc, 'Packing List EN', 'packing_list')

        # ── Proforma Invoice ──────────────────────────────────────────────────
        doc = Document()
        h = doc.add_heading('PROFORMA INVOICE', 0)
        h.alignment = A.CENTER
        doc.add_paragraph('')

        for label, var in [
            ('Invoice №:', '{{invoice_number}}'),
            ('Date:', '{{invoice_date}}'),
            ('Due Date:', '{{due_date}}'),
        ]:
            p = doc.add_paragraph()
            p.add_run(f'{label} ').bold = True
            p.add_run(var)

        doc.add_paragraph('')
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        tbl.rows[0].cells[0].text = (
            'SELLER:\n{{shipper_name}}\n{{shipper_address}}\n'
            'VAT: {{vat_number}}\nEORI: {{eori_number}}\n'
            'IBAN: {{bank_iban}}\nSWIFT: {{bank_swift}}')
        tbl.rows[0].cells[1].text = (
            'BUYER:\n{{customer_name}}\n{{customer_address}}\n'
            '{{customer_city}}, {{customer_country}}\n'
            'Email: {{customer_email}}\nVAT: {{customer_vat}}')

        doc.add_paragraph('')
        doc.add_heading('Items', level=2)
        tbl2 = doc.add_table(rows=2, cols=5)
        tbl2.style = 'Table Grid'
        for i, h in enumerate(['SKU', 'Description', 'Qty', 'Unit Price', 'Amount']):
            c = tbl2.rows[0].cells[i]
            c.text = h
            c.paragraphs[0].runs[0].bold = True
        r = tbl2.rows[1]
        r.cells[0].text = '{%tr for item in items %}{{item.sku}}'
        r.cells[1].text = '{{item.name}}'
        r.cells[2].text = '{{item.quantity}}'
        r.cells[3].text = '{{item.unit_price}} {{currency}}'
        r.cells[4].text = '{{item.total_price}} {{currency}}{%tr endfor %}'

        doc.add_paragraph('')
        tbl3 = doc.add_table(rows=3, cols=2)
        tbl3.style = 'Table Grid'
        for i, (label, val) in enumerate([
            ('Subtotal:', '{{subtotal}} {{currency}}'),
            ('VAT ({{vat_rate}}%):', '{{vat_amount}} {{currency}}'),
            ('TOTAL:', '{{total_amount}} {{currency}}'),
        ]):
            tbl3.rows[i].cells[0].text = label
            tbl3.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            tbl3.rows[i].cells[1].text = val

        doc.add_paragraph('')
        p = doc.add_paragraph()
        p.add_run('Payment Terms: ').bold = True
        p.add_run('{{payment_terms}}')
        doc.add_paragraph('{{proforma_notes}}')
        doc.add_paragraph('Generated: {{generated_date}}')
        save(doc, 'Proforma Invoice EN', 'proforma')

        # ── CN23 ──────────────────────────────────────────────────────────────
        doc = Document()
        h = doc.add_heading('CN23 CUSTOMS DECLARATION', 0)
        h.alignment = A.CENTER
        doc.add_paragraph('(For commercial goods sent by post)')
        doc.add_paragraph('')

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        tbl.rows[0].cells[0].text = (
            'SENDER:\n{{shipper_name}}\n{{shipper_address}}\n'
            '{{shipper_city}}, {{shipper_country}}\n{{shipper_email}}')
        tbl.rows[0].cells[1].text = (
            'ADDRESSEE:\n{{customer_name}}\n{{customer_address}}\n'
            '{{customer_city}}, {{customer_country}}\n{{customer_email}}')

        doc.add_paragraph('')
        for label, var in [
            ('Type of contents:', '{{customs_type}} — {{customs_reason}}'),
            ('Country of origin:', '{{country_of_origin}}'),
        ]:
            p = doc.add_paragraph()
            p.add_run(f'{label} ').bold = True
            p.add_run(var)

        doc.add_paragraph('')
        doc.add_heading('Contents', level=2)
        tbl2 = doc.add_table(rows=2, cols=5)
        tbl2.style = 'Table Grid'
        for i, h in enumerate(['HS Code', 'Description', 'Qty', 'Value', 'Weight']):
            c = tbl2.rows[0].cells[i]
            c.text = h
            c.paragraphs[0].runs[0].bold = True
        r = tbl2.rows[1]
        r.cells[0].text = '{%tr for item in items %}{{item.hs_code}}'
        r.cells[1].text = '{{item.name}}'
        r.cells[2].text = '{{item.quantity}}'
        r.cells[3].text = '{{item.total_price}} {{currency}}'
        r.cells[4].text = '{{item.weight}} kg{%tr endfor %}'

        doc.add_paragraph('')
        tbl3 = doc.add_table(rows=2, cols=2)
        tbl3.style = 'Table Grid'
        tbl3.rows[0].cells[0].text = 'Total declared value:'
        tbl3.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        tbl3.rows[0].cells[1].text = '{{declared_value}} {{currency}}'
        tbl3.rows[1].cells[0].text = 'Gross weight:'
        tbl3.rows[1].cells[0].paragraphs[0].runs[0].bold = True
        tbl3.rows[1].cells[1].text = '{{gross_weight}} kg'

        doc.add_paragraph('')
        doc.add_paragraph(
            'I certify that the particulars given in this declaration '
            'are correct and that this item does not contain any dangerous '
            'article or articles prohibited by legislation or by postal or '
            'customs regulations.')
        doc.add_paragraph('')
        p = doc.add_paragraph()
        p.add_run('Signature: ').bold = True
        p.add_run('_________________________')
        p = doc.add_paragraph()
        p.add_run('Date: ').bold = True
        p.add_run('{{invoice_date}}')
        doc.add_paragraph('Generated by Minerva BI: {{generated_date}}')
        save(doc, 'CN23 Customs Declaration EN', 'cn23')

        self.stdout.write(self.style.SUCCESS('Done. Templates ready!'))
        self.stdout.write('  See: /admin/documents/documenttemplate/')
